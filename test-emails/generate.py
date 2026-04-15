"""Generate UAT test attachments for Linkworks CRM."""
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import os

OUT = os.path.dirname(os.path.abspath(__file__))
styles = getSampleStyleSheet()
h = styles["Heading2"]
n = styles["Normal"]


def test1_pdf():
    doc = SimpleDocTemplate(os.path.join(OUT, "test1_standard_booking.pdf"), pagesize=A4)
    story = [
        Paragraph("Booking Request", h),
        Spacer(1, 12),
        Paragraph("Dear Linkworks team,", n),
        Spacer(1, 8),
        Paragraph("Please arrange a standard delivery for the shipment below.", n),
        Spacer(1, 12),
        Paragraph("<b>Pickup Address:</b> Unit 14, Acorn Industrial Estate, Birmingham B11 2PL, UK", n),
        Paragraph("<b>Delivery Address:</b> 27 Riverside Park, Manchester M17 1SN, UK", n),
        Spacer(1, 8),
        Paragraph("<b>Pickup Date:</b> 18 April 2026, 09:00 - 11:00", n),
        Paragraph("<b>Delivery Date:</b> 18 April 2026, by 17:00", n),
        Spacer(1, 8),
        Paragraph("<b>Cargo:</b> 6 pallets of packaged textiles", n),
        Paragraph("<b>Total Weight:</b> 850 kg", n),
        Paragraph("<b>Dimensions:</b> standard Euro pallets (1.2m x 0.8m x 1.5m)", n),
        Paragraph("<b>Hazardous:</b> No", n),
        Paragraph("<b>Vehicle Required:</b> Standard box van (7.5t)", n),
        Spacer(1, 12),
        Paragraph("Kind regards,<br/>James Harper<br/>Harper Textiles Ltd", n),
    ]
    doc.build(story)


def test3_pdf():
    doc = SimpleDocTemplate(os.path.join(OUT, "test3_multi_order_booking.pdf"), pagesize=A4)
    story = [
        Paragraph("Bulk Booking Request - 5 Consignments", h),
        Spacer(1, 12),
        Paragraph("Please arrange the following deliveries for this week. Pricing confirmation needed by EOD.", n),
        Spacer(1, 12),
    ]

    orders = [
        ["#", "Pickup", "Delivery", "Cargo", "Weight", "Hazardous", "Vehicle"],
        ["1", "Leeds LS1 3AA", "Glasgow G2 5AD",
         "Industrial solvents (UN 1993, Cl 3)", "2,400 kg", "Yes - Class 3 flammable",
         "Standard with ADR certification"],
        ["2", "Dover CT16 1JA", "Bristol BS1 6QH",
         "Lithium-ion battery packs (UN 3480)", "1,850 kg", "Yes - Class 9",
         "Standard with ADR"],
        ["3", "Felixstowe IP11 3TW", "Nottingham NG7 2RU",
         "Steel coils", "18,500 kg", "No", "Heavy goods (26t articulated)"],
        ["4", "Southampton SO15 1GH", "Leeds LS11 5DL",
         "Granite slabs, crated", "14,200 kg", "No", "Flatbed curtain-sided (18t)"],
        ["5", "London E16 4QH", "Cardiff CF10 1EP",
         "Office furniture, flat-packed", "620 kg", "No", "Standard box van"],
    ]
    table = Table(orders, colWidths=[0.7*cm, 3.2*cm, 3.2*cm, 3.8*cm, 2.0*cm, 2.8*cm, 3.8*cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(table)
    story.append(Spacer(1, 16))
    story.append(Paragraph("Preferred pickup window: 20-22 April 2026.", n))
    story.append(Paragraph("All deliveries tail-lift not required unless noted.", n))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Regards,<br/>Priya Shah<br/>Logistics Manager, Apex Freight Solutions", n))
    doc.build(story)


def test4_docx():
    doc = Document()
    title = doc.add_heading("Out-of-Gauge Booking Request", level=1)
    doc.add_paragraph("Hi team,")
    doc.add_paragraph(
        "We need to book an out-of-gauge shipment for a client in the renewables sector. "
        "Please confirm availability and pricing at the earliest."
    )

    doc.add_heading("Shipment Details", level=2)
    details = [
        ("Pickup Address", "Siemens Gamesa Facility, Green Port Hull, Hedon Road, Hull HU9 5PQ, UK"),
        ("Delivery Address", "Teesside Offshore Assembly Yard, South Gare, Redcar TS10 5QW, UK"),
        ("Pickup Date", "24 April 2026, 06:00 (access via booked slot only)"),
        ("Delivery Date", "25 April 2026, by 18:00"),
        ("Cargo", "Wind turbine blade section, 1 unit"),
        ("Dimensions", "Length 58.5 m, Width 3.8 m, Height 3.2 m"),
        ("Weight", "24,500 kg"),
        ("Hazardous", "No"),
        ("Vehicle Required", "Out-of-gauge abnormal load trailer with steerable rear axle"),
        ("Escort Required", "Yes - 2x private escort vehicles, police notification already filed (ref PN-NE-2604-18)"),
    ]
    for k, v in details:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)

    doc.add_paragraph("")
    doc.add_paragraph("Route survey document can be provided on request. Please reply with a quote.")
    doc.add_paragraph("")
    doc.add_paragraph("Thanks,")
    doc.add_paragraph("Martin Ellis")
    doc.add_paragraph("Projects Coordinator, NorthSea Renewables Ltd")
    doc.add_paragraph("m.ellis@northsea-renewables.example")

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.size:
                run.font.size = Pt(11)

    doc.save(os.path.join(OUT, "test4_oog_booking.docx"))


def test6_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Bookings"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="305496")
    centre = Alignment(horizontal="center", vertical="center", wrap_text=True)

    headers = [
        "Order Ref (internal)", "Pickup Address", "Delivery Address",
        "Pickup Date", "Cargo Description", "Weight (kg)",
        "Hazardous", "Vehicle Type", "Notes",
    ]
    for col, val in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=val)
        c.font = header_font
        c.fill = header_fill
        c.alignment = centre

    rows = [
        ["INT-1041", "Unit 5, Park Royal, London NW10 7NA",
         "22 Queens Road, Reading RG1 4BS",
         "22 April 2026, 08:00", "Pharmaceutical packaging supplies",
         480, "No", "Standard van",
         "Regular client, gate code 4421"],
        ["INT-1042", "Tilbury Docks, Essex RM18 7EH",
         "Logistics Hub, Doncaster DN4 5JW",
         "22 April 2026, 14:00", "Imported ceramic tiles on pallets",
         6800, "No", "Standard rigid (18t)",
         "Fork-lift available on site"],
        ["INT-1043", "BASF Site, Seal Sands, Middlesbrough TS2 1UB",
         "Industrial Park, Warrington WA5 1QR",
         "23 April 2026, 09:30", "Industrial cleaning chemicals (UN 1760)",
         1200, "Yes - Class 8 corrosive", "Standard with ADR driver",
         "Safety data sheet will be provided on collection"],
        ["INT-1044", "Farm supply depot, Norwich NR6 5DU",
         "Agri Co-op, Inverness IV1 1NH",
         "24 April 2026, 06:00", "Fertiliser bags, palletised",
         22000, "No", "Articulated curtain-sided (44t)",
         "Long haul, overnight run acceptable"],
    ]
    for r_idx, row in enumerate(rows, 2):
        for c_idx, val in enumerate(row, 1):
            ws.cell(row=r_idx, column=c_idx, value=val)

    widths = [16, 36, 36, 20, 32, 12, 22, 28, 34]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    wb.save(os.path.join(OUT, "test6_bulk_bookings.xlsx"))


def test7_pdf_tailift():
    doc = SimpleDocTemplate(os.path.join(OUT, "test7_tailift_booking.pdf"), pagesize=A4)
    story = [
        Paragraph("Delivery Booking - Tail-Lift Required", h),
        Spacer(1, 12),
        Paragraph("Hello,", n),
        Spacer(1, 8),
        Paragraph(
            "Please arrange collection and delivery of the goods below. "
            "Delivery site has no forklift so a tail-lift vehicle is essential.",
            n,
        ),
        Spacer(1, 12),
        Paragraph("<b>Customer Ref:</b> PO-45872-A", n),
        Paragraph("<b>Pickup:</b> Warehouse 7, Oakdale Business Park, Coventry CV3 4LF", n),
        Paragraph("<b>Delivery:</b> Riverside Cafe, 4 Market Square, Stratford-upon-Avon CV37 6AB", n),
        Spacer(1, 8),
        Paragraph("<b>Pickup Date:</b> 23 April 2026, 10:00 - 12:00", n),
        Paragraph("<b>Delivery Date:</b> Same day, by 16:00", n),
        Spacer(1, 8),
        Paragraph("<b>Cargo:</b> 4 x commercial refrigeration units on pallets", n),
        Paragraph("<b>Weight:</b> 1,150 kg total", n),
        Paragraph("<b>Dimensions:</b> each unit 1.8m H x 0.9m W x 0.7m D", n),
        Paragraph("<b>Hazardous:</b> No", n),
        Paragraph("<b>Vehicle:</b> 7.5t with tail-lift (mandatory - no forklift at delivery)", n),
        Spacer(1, 12),
        Paragraph("Regards,<br/>Owen Walsh<br/>CoolChain Supplies", n),
    ]
    doc.build(story)


if __name__ == "__main__":
    test1_pdf()
    test3_pdf()
    test4_docx()
    test6_xlsx()
    test7_pdf_tailift()
    print("Generated:")
    for f in sorted(os.listdir(OUT)):
        if f != "generate.py":
            print("  -", f)
